from __future__ import annotations

import re
from copy import deepcopy
from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

from core.template_engine import render_resume_html, render_resume_markdown


def build_export_bundle(resume: dict, profile: dict, template: dict) -> dict[str, bytes | str]:
    clean_resume = _sanitize_nested(resume)
    clean_profile = _sanitize_nested(profile)
    clean_template = _sanitize_nested(template)
    html = render_resume_html(resume=clean_resume, template=clean_template, profile=clean_profile)
    markdown = render_resume_markdown(resume=clean_resume, profile=clean_profile)
    return {
        "html": html,
        "markdown": markdown,
        "docx": build_docx_bytes(resume=clean_resume, profile=clean_profile, template=clean_template),
        "pdf": build_pdf_bytes(resume=clean_resume, profile=clean_profile, template=clean_template),
    }


def build_docx_bytes(resume: dict, profile: dict, template: dict) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    font_name = _docx_font_name(template.get("font_family", "Arial, sans-serif"))
    base_color = _hex_to_rgb(template.get("text_color", "#111827"))
    accent_color = _hex_to_rgb(template.get("accent_color", "#1f4e79"))
    muted_color = _hex_to_rgb(template.get("muted_color", "#4b5563"))
    emphasis_keywords = _effective_bold_keywords(resume)

    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal_style.font.size = Pt(10.2)
    normal_style.font.color.rgb = RGBColor(*base_color)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(name, before=0, after=1)
    run = name.add_run(profile.get("name", ""))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(*base_color)
    run.font.name = font_name

    headline = document.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(headline, before=0, after=1)
    _append_text_with_bold_keywords(
        headline,
        resume.get("headline", ""),
        emphasis_keywords,
        font_name=font_name,
        font_size=11.4,
        color=accent_color,
        bold=True,
    )

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(meta, before=0, after=7)
    meta_run = meta.add_run(_meta_line(profile))
    meta_run.font.size = Pt(9.4)
    meta_run.font.color.rgb = RGBColor(*muted_color)
    meta_run.font.name = font_name

    if template.get("header_style", "rule") == "rule":
        border = document.add_paragraph()
        _set_bottom_border(border, color=template.get("accent_color", "#1f4e79"), size=10)
        _set_spacing(border, before=0, after=8)

    section_order = template.get(
        "section_order",
        ["summary", "technical_skills", "work_history", "education_history"],
    )
    for section_key in section_order:
        if section_key == "summary":
            _add_heading(document, "PROFESSIONAL SUMMARY", font_name, accent_color)
            _add_body_paragraph(document, resume.get("summary", ""), font_name, base_color, emphasis_keywords)
        elif section_key == "technical_skills":
            _add_heading(document, "TECHNICAL SKILLS", font_name, accent_color)
            skill_style = template.get("skill_style", "grouped_bullets")
            if skill_style in {"grouped", "grouped_bullets"}:
                for group in _resolve_skill_groups(resume):
                    _add_bullet_with_bold_label(
                        document=document,
                        label=group.get("category", "Other Relevant"),
                        value=", ".join(group.get("items", [])),
                        font_name=font_name,
                        base_color=base_color,
                        emphasis_keywords=emphasis_keywords,
                    )
            else:
                _add_body_paragraph(
                    document,
                    _skill_line(resume.get("technical_skills", []), skill_style),
                    font_name,
                    base_color,
                    emphasis_keywords,
                )
        elif section_key == "work_history":
            _add_heading(document, "PROFESSIONAL EXPERIENCE", font_name, accent_color)
            for job in resume.get("work_history", []):
                _add_job_block(document, job, template, font_name, base_color, muted_color, emphasis_keywords)
        elif section_key == "education_history":
            _add_heading(document, "EDUCATION", font_name, accent_color)
            for item in resume.get("education_history", []):
                p = document.add_paragraph()
                _set_spacing(p, before=0, after=0)
                r1 = p.add_run(item.get("university", ""))
                r1.bold = True
                r1.font.name = font_name
                r1.font.size = Pt(10.3)
                r1.font.color.rgb = RGBColor(*base_color)

                degree_paragraph = document.add_paragraph()
                degree_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _set_spacing(degree_paragraph, before=0, after=0)
                _append_text_with_bold_keywords(
                    degree_paragraph,
                    item.get("degree", ""),
                    emphasis_keywords,
                    font_name=font_name,
                    font_size=10.0,
                    color=base_color,
                )
                meta_line = " | ".join(part for part in [item.get("duration", ""), item.get("location", "")] if part)
                if meta_line:
                    meta_paragraph = document.add_paragraph()
                    _set_spacing(meta_paragraph, before=0, after=4)
                    r2 = meta_paragraph.add_run(meta_line)
                    r2.italic = True
                    r2.font.name = font_name
                    r2.font.size = Pt(9.2)
                    r2.font.color.rgb = RGBColor(*muted_color)

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_pdf_bytes(resume: dict, profile: dict, template: dict) -> bytes:
    stream = BytesIO()
    doc = SimpleDocTemplate(
        stream,
        pagesize=LETTER,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )

    palette = {
        "accent": _hex_to_reportlab(template.get("accent_color", "#1f4e79")),
        "text": _hex_to_reportlab(template.get("text_color", "#111827")),
        "muted": _hex_to_reportlab(template.get("muted_color", "#4b5563")),
    }
    font_name = _pdf_font_name(template.get("font_family", "Arial, sans-serif"))
    styles = getSampleStyleSheet()
    body_size = 9.8 if template.get("density", "normal") == "tight" else 10.2
    leading = 12.3 if template.get("density", "normal") == "tight" else 13.1
    emphasis_keywords = _effective_bold_keywords(resume)

    styles.add(ParagraphStyle(name="ResumeName", parent=styles["Normal"], fontName=font_name, fontSize=19.5, leading=22, alignment=TA_CENTER, textColor=palette["text"], spaceAfter=1))
    styles.add(ParagraphStyle(name="ResumeHeadline", parent=styles["Normal"], fontName=font_name, fontSize=11.4, leading=14, alignment=TA_CENTER, textColor=palette["accent"], spaceAfter=2))
    styles.add(ParagraphStyle(name="ResumeMeta", parent=styles["Normal"], fontName=font_name, fontSize=8.9, leading=11, alignment=TA_CENTER, textColor=palette["muted"], spaceAfter=8))
    styles.add(ParagraphStyle(name="SectionTitle", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=12, textColor=palette["accent"], spaceBefore=8, spaceAfter=4, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="BodyResume", parent=styles["Normal"], fontName=font_name, fontSize=body_size, leading=leading, textColor=palette["text"], spaceAfter=2, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name="JobCompanyLine", parent=styles["Normal"], fontName=font_name, fontSize=10.3, leading=12.6, textColor=palette["text"], alignment=TA_LEFT, spaceAfter=1))
    styles.add(ParagraphStyle(name="JobRole", parent=styles["Normal"], fontName=font_name, fontSize=10.0, leading=12.1, textColor=palette["text"], spaceAfter=1, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name="JobMeta", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=11, textColor=palette["muted"], spaceAfter=1, alignment=TA_JUSTIFY))

    story = [
        Paragraph(_escape(profile.get("name", "")), styles["ResumeName"]),
        Paragraph(_pdf_markup(resume.get("headline", ""), emphasis_keywords, force_bold=True), styles["ResumeHeadline"]),
        Paragraph(_escape(_meta_line(profile)), styles["ResumeMeta"]),
    ]

    if template.get("header_style", "rule") == "rule":
        story.append(HRFlowable(width="100%", thickness=1.2, color=palette["accent"], spaceBefore=0, spaceAfter=8))

    section_order = template.get(
        "section_order",
        ["summary", "technical_skills", "work_history", "education_history"],
    )
    for section_key in section_order:
        if section_key == "summary":
            story.append(Paragraph("PROFESSIONAL SUMMARY", styles["SectionTitle"]))
            story.append(Paragraph(_pdf_markup(resume.get("summary", ""), emphasis_keywords), styles["BodyResume"]))
        elif section_key == "technical_skills":
            story.append(Paragraph("TECHNICAL SKILLS", styles["SectionTitle"]))
            skill_style = template.get("skill_style", "grouped_bullets")
            if skill_style in {"grouped", "grouped_bullets"}:
                bullet_items = [
                    ListItem(
                        Paragraph(
                            f"<b>{_escape(group.get('category', 'Other Relevant'))}:</b> {_pdf_markup(', '.join(group.get('items', [])), emphasis_keywords)}",
                            styles["BodyResume"],
                        )
                    )
                    for group in _resolve_skill_groups(resume)
                    if group.get("items")
                ]
                if bullet_items:
                    story.append(ListFlowable(bullet_items, bulletType="bullet", leftIndent=14))
            else:
                story.append(Paragraph(_pdf_markup(_skill_line(resume.get("technical_skills", []), skill_style), emphasis_keywords), styles["BodyResume"]))
        elif section_key == "work_history":
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", styles["SectionTitle"]))
            for job in resume.get("work_history", []):
                meta_inline = " | ".join(part for part in [job.get("duration", ""), job.get("location", "")] if part)
                company_html = f"<b>{_escape(job.get('company_name', ''))}</b>"
                if meta_inline:
                    company_html += f" <font color='{template.get('muted_color', '#4b5563')}'>| {_escape(meta_inline)}</font>"
                story.append(Paragraph(company_html, styles["JobCompanyLine"]))
                story.append(Paragraph(_pdf_markup(job.get("role_title", ""), emphasis_keywords), styles["JobRole"]))
                if template.get("show_role_headline", True) and job.get("role_headline"):
                    story.append(Paragraph(f"<i>{_pdf_markup(job.get('role_headline', ''), emphasis_keywords)}</i>", styles["JobMeta"]))
                bullet_items = [ListItem(Paragraph(_pdf_markup(bullet, emphasis_keywords), styles["BodyResume"])) for bullet in job.get("bullets", [])]
                if bullet_items:
                    story.append(ListFlowable(bullet_items, bulletType="bullet", leftIndent=14))
                story.append(Spacer(1, 5))
        elif section_key == "education_history":
            story.append(Paragraph("EDUCATION", styles["SectionTitle"]))
            for item in resume.get("education_history", []):
                edu_html = f"<b>{_escape(item.get('university', ''))}</b><br/>{_pdf_markup(item.get('degree', ''), emphasis_keywords)}"
                meta_line = " | ".join(part for part in [item.get("duration", ""), item.get("location", "")] if part)
                if meta_line:
                    edu_html += f"<br/><font color='{template.get('muted_color', '#4b5563')}'>{_escape(meta_line)}</font>"
                story.append(Paragraph(edu_html, styles["BodyResume"]))

    doc.build(story)
    return stream.getvalue()


def _add_heading(document: Document, title: str, font_name: str, accent_color: tuple[int, int, int]) -> None:
    p = document.add_paragraph()
    _set_spacing(p, before=8, after=3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.3)
    run.font.name = font_name
    run.font.color.rgb = RGBColor(*accent_color)


def _add_body_paragraph(document: Document, text: str, font_name: str, base_color: tuple[int, int, int], emphasis_keywords: list[str]) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, before=0, after=3)
    _append_text_with_bold_keywords(
        p,
        text,
        emphasis_keywords,
        font_name=font_name,
        font_size=10.0,
        color=base_color,
    )


def _add_bullet_with_bold_label(document: Document, label: str, value: str, font_name: str, base_color: tuple[int, int, int], emphasis_keywords: list[str]) -> None:
    paragraph = document.add_paragraph(style=document.styles["Normal"])
    _set_spacing(paragraph, before=0, after=0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.15)
    bullet_run = paragraph.add_run("• ")
    bullet_run.font.name = font_name
    bullet_run.font.size = Pt(10.0)
    bullet_run.font.color.rgb = RGBColor(*base_color)

    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.name = font_name
    label_run.font.size = Pt(10.0)
    label_run.font.color.rgb = RGBColor(*base_color)

    _append_text_with_bold_keywords(
        paragraph,
        value,
        emphasis_keywords,
        font_name=font_name,
        font_size=10.0,
        color=base_color,
    )


def _add_job_block(document: Document, job: dict, template: dict, font_name: str, base_color: tuple[int, int, int], muted_color: tuple[int, int, int], emphasis_keywords: list[str]) -> None:
    header = document.add_paragraph()
    _set_spacing(header, before=0, after=0)
    company_run = header.add_run(job.get("company_name", ""))
    company_run.bold = True
    company_run.font.name = font_name
    company_run.font.size = Pt(10.5)
    company_run.font.color.rgb = RGBColor(*base_color)

    meta_inline = " | ".join(part for part in [job.get("duration", ""), job.get("location", "")] if part)
    if meta_inline:
        meta_run = header.add_run(f" | {meta_inline}")
        meta_run.font.name = font_name
        meta_run.font.size = Pt(9.1)
        meta_run.font.color.rgb = RGBColor(*muted_color)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(role, before=0, after=0)
    _append_text_with_bold_keywords(
        role,
        job.get("role_title", ""),
        emphasis_keywords,
        font_name=font_name,
        font_size=10.0,
        color=base_color,
        bold=True,
    )

    if template.get("show_role_headline", True) and job.get("role_headline"):
        role_headline = document.add_paragraph()
        role_headline.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(role_headline, before=0, after=0)
        _append_text_with_bold_keywords(
            role_headline,
            job.get("role_headline", ""),
            emphasis_keywords,
            font_name=font_name,
            font_size=9.2,
            color=muted_color,
            italic=True,
        )

    for bullet in job.get("bullets", []):
        bp = document.add_paragraph(style=document.styles["Normal"])
        _set_spacing(bp, before=0, after=0)
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.left_indent = Inches(0.18)
        bp.paragraph_format.first_line_indent = Inches(-0.15)
        bullet_prefix = bp.add_run("• ")
        bullet_prefix.font.name = font_name
        bullet_prefix.font.size = Pt(10.0)
        bullet_prefix.font.color.rgb = RGBColor(*base_color)
        _append_text_with_bold_keywords(
            bp,
            bullet,
            emphasis_keywords,
            font_name=font_name,
            font_size=10.0,
            color=base_color,
        )

    spacer = document.add_paragraph()
    _set_spacing(spacer, before=0, after=4)


def _append_text_with_bold_keywords(paragraph, text: str, emphasis_keywords: list[str], font_name: str, font_size: float, color: tuple[int, int, int], bold: bool = False, italic: bool = False) -> None:
    for segment_text, matched in _split_by_keywords(text, emphasis_keywords):
        if not segment_text:
            continue
        run = paragraph.add_run(segment_text)
        run.bold = bold or matched
        run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*color)


def _resolve_skill_groups(resume: dict) -> list[dict]:
    groups = resume.get("skill_groups") or []
    normalized: list[dict] = []
    for group in groups:
        category = str(group.get("category", "")).strip()
        items = [str(item).strip() for item in group.get("items", []) if str(item).strip()]
        if category and items:
            normalized.append({"category": category, "items": items})
    if normalized:
        return normalized
    skills = [str(item).strip() for item in resume.get("technical_skills", []) if str(item).strip()]
    return [{"category": "Other Relevant", "items": skills}] if skills else []


MULTI_WORD_TECH_PHRASES = [
    "React Native",
    "React Hooks",
    "React Router",
    "React Query",
    "React Testing Library",
    "Redux Toolkit",
    "Redux Saga",
    "Redux Thunk",
    "Vue Native",
    "Vue Router",
    "Vue Test Utils",
    "Spring Boot",
    "Spring Cloud",
    "Spring Security",
    "Spring Data",
    "Apollo Client",
    "Apollo Server",
    "Apollo GraphQL",
    "Material UI",
    "Tailwind CSS",
    "Next.js",
    "Nuxt.js",
    "Node.js",
    "Express.js",
    "Nest.js",
    "Vue.js",
    "Ember.js",
    "Backbone.js",
    "Three.js",
    "D3.js",
    "Socket.IO",
    "Ruby on Rails",
    "ASP.NET Core",
    "Entity Framework",
    "SQL Server",
    "Azure DevOps",
    "Azure Functions",
    "Cloud Run",
    "GitHub Actions",
    "GitLab CI",
    "Argo CD",
    "Hugging Face",
    "OpenAI API",
    "New Relic",
    "Key Vault",
    "Web Components",
    "Service Worker",
    "Single Page Application",
    "Server Side Rendering",
    "Client Side Rendering",
]


def _expanded_keywords_with_phrases(keywords: list[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for item in keywords or []:
        clean = str(item).strip()
        if not clean:
            continue
        key = clean.lower()
        if key in seen:
            continue
        seen.add(key)
        ordered.append(clean)
    base_lower = {kw.lower(): kw for kw in ordered}
    for phrase in MULTI_WORD_TECH_PHRASES:
        first_token = phrase.split()[0].lower()
        if first_token in base_lower and phrase.lower() not in seen:
            seen.add(phrase.lower())
            ordered.append(phrase)
    return ordered


def _effective_bold_keywords(resume: dict) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()
    manual = resume.get("bold_keywords", []) or []
    auto_keywords = resume.get("fit_keywords", []) if resume.get("auto_bold_fit_keywords", True) else []
    technical_skills = resume.get("technical_skills", []) or []
    for item in [*manual, *auto_keywords, *technical_skills]:
        clean = str(item).strip()
        if not clean:
            continue
        key = clean.lower()
        if key not in seen:
            seen.add(key)
            ordered.append(clean)
    return _expanded_keywords_with_phrases(ordered)


def _split_by_keywords(text: str, keywords: list[str]) -> list[tuple[str, bool]]:
    if not text:
        return []
    if not keywords:
        return [(text, False)]
    pattern = _keyword_pattern(keywords)
    parts: list[tuple[str, bool]] = []
    last = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last:
            parts.append((text[last:start], False))
        parts.append((text[start:end], True))
        last = end
    if last < len(text):
        parts.append((text[last:], False))
    return parts or [(text, False)]


def _keyword_pattern(keywords: list[str]) -> re.Pattern[str]:
    expanded = _expanded_keywords_with_phrases(list(keywords or []))
    ordered = sorted({str(item).strip() for item in expanded if str(item).strip()}, key=len, reverse=True)
    if not ordered:
        return re.compile(r"(?!x)x")
    escaped_terms = [re.escape(item) for item in ordered]
    return re.compile(rf"(?<![A-Za-z0-9])(?:{'|'.join(escaped_terms)})(?![A-Za-z0-9])", re.IGNORECASE)


def _docx_font_name(font_family: str) -> str:
    lower = font_family.lower()
    if "times" in lower or "georgia" in lower:
        return "Times New Roman"
    if "calibri" in lower:
        return "Calibri"
    return "Arial"


def _pdf_font_name(font_family: str) -> str:
    lower = font_family.lower()
    if "times" in lower or "georgia" in lower:
        return "Times-Roman"
    if "courier" in lower:
        return "Courier"
    return "Helvetica"


def _meta_line(profile: dict) -> str:
    return " | ".join(
        part
        for part in [
            profile.get("email", ""),
            profile.get("phone", ""),
            profile.get("location", ""),
            profile.get("linkedin", ""),
            profile.get("portfolio", ""),
        ]
        if part
    )


def _skill_line(skills: Iterable[str], skill_style: str) -> str:
    items = [item for item in skills if item]
    if skill_style == "pipe":
        return " | ".join(items)
    return ", ".join(items)


def _hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip().lstrip("#")
    if len(value) != 6:
        return (17, 24, 39)
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def _hex_to_reportlab(value: str) -> colors.Color:
    r, g, b = _hex_to_rgb(value)
    return colors.Color(r / 255.0, g / 255.0, b / 255.0)


def _set_spacing(paragraph, before: int = 0, after: int = 0, line: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def _set_bottom_border(paragraph, color: str = "1f4e79", size: int = 8) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.replace("#", ""))
    pbdr.append(bottom)
    pPr.append(pbdr)


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")


def _pdf_markup(text: str, keywords: list[str], force_bold: bool = False) -> str:
    segments = _split_by_keywords(text, keywords)
    if not segments:
        escaped = _escape(text)
        return f"<b>{escaped}</b>" if force_bold and escaped else escaped
    parts: list[str] = []
    for segment, matched in segments:
        escaped = _escape(segment)
        if force_bold or matched:
            parts.append(f"<b>{escaped}</b>")
        else:
            parts.append(escaped)
    return "".join(parts)


def _sanitize_nested(value):
    if isinstance(value, dict):
        return {key: _sanitize_nested(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_sanitize_nested(item) for item in value]
    if isinstance(value, str):
        return _sanitize_text(value)
    return deepcopy(value)


def _sanitize_text(text: str) -> str:
    replacements = {
        "‐": "-",
        "‑": "-",
        "‒": "-",
        "–": "-",
        "—": "-",
        "−": "-",
        "‘": "'",
        "’": "'",
        '“': '"',
        '”': '"',
        "…": "...",
        " ": " ",
    }
    cleaned = text
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    return " ".join(cleaned.split())
