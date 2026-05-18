from __future__ import annotations

import base64
import html
import json
import re
import shutil
import subprocess
import sys
import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.resume_engine import KNOWN_TECH_TERMS


SECTION_ALIASES = {
    "summary": {
        "summary",
        "professional summary",
        "profile",
        "about",
        "about me",
        "professional profile",
        "profile summary",
        "career profile",
        "career summary",
    },
    "skills": {
        "skills",
        "technical skills",
        "core skills",
        "core competencies",
        "technologies",
        "technology stack",
    },
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment history",
        "work history",
    },
    "education": {
        "education",
        "education history",
        "academic background",
    },
    "projects": {"projects", "selected projects"},
    "certifications": {"certifications", "certificates", "licenses"},
}
ALL_SECTION_TITLES = {title for titles in SECTION_ALIASES.values() for title in titles}
TITLE_PLACEHOLDERS = {"___resume_title___", "___headline___", "__resume_title__", "__headline__"}
EXP_ROLE_PLACEHOLDERS = {"___title___", "__role__"}
SUMMARY_PLACEHOLDERS = {"___summary___", "___professional_summary___"}
SKILL_PLACEHOLDERS = {"___skills___", "___technical_skills___"}
EXPERIENCE_PLACEHOLDERS = {"___experience___", "___work_experience___", "___professional_experience___"}

_DATE_RE = re.compile(
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}\b|\b\d{4}\s*[–—-]\s*(?:present|current|\d{4})\b",
    re.IGNORECASE,
)
_ROLE_MARKER_PATTERN = re.compile(r"(___title___|__role__)", re.IGNORECASE)


def _clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _strip_markdown_emphasis(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(?![a-z_]+__)(.*?)__", r"\1", text)
    return text


def _plain_resume_text(value: object) -> str:
    return _clean_text(_strip_markdown_emphasis(value))


def _preserve_line_text(value: object) -> str:
    return _strip_markdown_emphasis(value).replace("\r", "").replace("\n", " ")


def _paragraph_contains_any_marker(paragraph: Paragraph, markers: set[str]) -> bool:
    text_lower = (paragraph.text or "").lower()
    return any(marker.lower() in text_lower for marker in markers)


def _paragraphs_in_container(container) -> Iterable[Paragraph]:
    """Yield paragraphs in real visual DOCX order, including tables.

    Some resumes use one-cell tables only for section headers. Using
    container.paragraphs first and container.tables second makes section ranges
    wrong. Walking the OOXML children preserves the visible order:
    paragraph -> table heading -> following outside-table content.
    """
    if hasattr(container, "_body"):
        parent_elm = getattr(container._body, "_element", None)
    else:
        parent_elm = getattr(container, "_element", None) or getattr(container, "element", None)

    if parent_elm is None:
        for paragraph in getattr(container, "paragraphs", []) or []:
            yield paragraph
        for table in getattr(container, "tables", []) or []:
            for row in table.rows:
                for cell in row.cells:
                    yield from _paragraphs_in_container(cell)
        return

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, container)
        elif child.tag == qn("w:tbl"):
            table = Table(child, container)
            seen_cells: set[int] = set()
            for row in table.rows:
                for cell in row.cells:
                    key = id(cell._tc)
                    if key in seen_cells:
                        continue
                    seen_cells.add(key)
                    yield from _paragraphs_in_container(cell)


def _all_story_paragraphs(doc: Document) -> list[Paragraph]:
    paragraphs = list(_paragraphs_in_container(doc))
    seen_parts = {id(doc.part)}
    for section in doc.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            try:
                part_id = id(container.part)
            except Exception:
                continue
            if part_id in seen_parts:
                continue
            seen_parts.add(part_id)
            paragraphs.extend(list(_paragraphs_in_container(container)))
    return paragraphs


def _all_body_paragraphs(doc: Document) -> list[Paragraph]:
    return list(_paragraphs_in_container(doc))


def _first_run_or_none(paragraph: Paragraph):
    return paragraph.runs[0] if paragraph.runs else None


def _first_meaningful_run_or_none(paragraph: Paragraph):
    for run in paragraph.runs:
        if str(run.text or "").strip():
            return run
    return _first_run_or_none(paragraph)


def _clear_paragraph_keep_ppr(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _copy_run_format(source_run, target_run) -> None:
    if source_run is None:
        return
    try:
        if source_run._r.rPr is not None:
            target_run._r.insert(0, deepcopy(source_run._r.rPr))
    except Exception:
        pass


def _force_run_not_bold(run) -> None:
    try:
        run.bold = False
    except Exception:
        pass
    try:
        rpr = run._r.get_or_add_rPr()
        for tag in ("w:b", "w:bCs"):
            node = rpr.find(qn(tag))
            if node is not None:
                rpr.remove(node)
        rstyle = rpr.find(qn("w:rStyle"))
        if rstyle is not None:
            style_id = str(rstyle.get(qn("w:val"), "")).strip().lower()
            if style_id in {"strong", "bold", "13"}:
                rpr.remove(rstyle)
    except Exception:
        pass


def _force_run_bold(run) -> None:
    try:
        run.bold = True
    except Exception:
        pass
    try:
        rpr = run._r.get_or_add_rPr()
        if rpr.find(qn("w:b")) is None:
            rpr.append(OxmlElement("w:b"))
        if rpr.find(qn("w:bCs")) is None:
            rpr.append(OxmlElement("w:bCs"))
    except Exception:
        pass


MULTI_WORD_TECH_PHRASES = [
    "React Native",
    "React Hooks",
    "React Router",
    "React Query",
    "React Testing Library",
    "Redux Toolkit",
    "Redux Saga",
    "Redux Thunk",
    "Vue Native",
    "Vue Router",
    "Vue Test Utils",
    "Spring Boot",
    "Spring Cloud",
    "Spring Security",
    "Spring Data",
    "Apollo Client",
    "Apollo Server",
    "Apollo GraphQL",
    "Material UI",
    "Tailwind CSS",
    "Next.js",
    "Nuxt.js",
    "Node.js",
    "Express.js",
    "Nest.js",
    "Vue.js",
    "Ember.js",
    "Backbone.js",
    "Three.js",
    "D3.js",
    "Socket.IO",
    "Ruby on Rails",
    "ASP.NET Core",
    "Entity Framework",
    "SQL Server",
    "Azure DevOps",
    "Azure Functions",
    "Cloud Run",
    "GitHub Actions",
    "GitLab CI",
    "Argo CD",
    "Hugging Face",
    "OpenAI API",
    "New Relic",
    "Key Vault",
    "Web Components",
    "Service Worker",
    "Single Page Application",
    "Server Side Rendering",
    "Client Side Rendering",
]


def _expanded_keywords_with_phrases(keywords: list[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for item in keywords or []:
        clean = str(item).strip()
        if not clean:
            continue
        key = clean.lower()
        if key in seen:
            continue
        seen.add(key)
        ordered.append(clean)
    base_lower = {kw.lower(): kw for kw in ordered}
    for phrase in MULTI_WORD_TECH_PHRASES:
        first_token = phrase.split()[0].lower()
        if first_token in base_lower and phrase.lower() not in seen:
            seen.add(phrase.lower())
            ordered.append(phrase)
    return ordered


def _technical_skill_keywords(resume: dict) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()
    # Resume-specific skills first (highest priority / preserves casing)
    for item in resume.get("technical_skills", []) or []:
        clean = _plain_resume_text(item)
        if len(clean) < 2:
            continue
        key = clean.lower()
        if key not in seen:
            seen.add(key)
            ordered.append(clean)
    # Add all known tech terms so bullets/summary always get fully bolded
    for term in KNOWN_TECH_TERMS:
        key = term.lower()
        if key not in seen:
            seen.add(key)
            ordered.append(term)
    return _expanded_keywords_with_phrases(ordered)


def _keyword_pattern(keywords: list[str]) -> re.Pattern[str]:
    expanded = _expanded_keywords_with_phrases(list(keywords or []))
    clean_keywords = sorted(
        {str(item).strip() for item in expanded if str(item).strip()},
        key=len,
        reverse=True,
    )
    if not clean_keywords:
        return re.compile(r"(?!x)x")
    parts: list[str] = []
    for keyword in clean_keywords:
        escaped = re.escape(keyword)
        # Avoid bolding inside larger words, while allowing punctuation in tech names.
        parts.append(rf"(?<![A-Za-z0-9]){escaped}(?![A-Za-z0-9])")
    return re.compile("|".join(parts), re.IGNORECASE)


def _add_run(
    paragraph: Paragraph,
    text: str,
    source_run=None,
    *,
    force_no_bold: bool = False,
    force_bold: bool = False,
):
    run = paragraph.add_run(text)
    _copy_run_format(source_run, run)
    if force_no_bold:
        _force_run_not_bold(run)
    if force_bold:
        _force_run_bold(run)
    return run


def _add_text_with_keyword_bold(
    paragraph: Paragraph,
    text: str,
    source_run=None,
    keywords: list[str] | None = None,
    *,
    base_no_bold: bool = True,
) -> None:
    text = _plain_resume_text(text)
    keywords = keywords or []
    if not keywords:
        _add_run(paragraph, text, source_run, force_no_bold=base_no_bold)
        return
    pattern = _keyword_pattern(keywords)
    pos = 0
    matched_any = False
    for match in pattern.finditer(text):
        if match.start() > pos:
            _add_run(paragraph, text[pos:match.start()], source_run, force_no_bold=base_no_bold)
        _add_run(paragraph, match.group(0), source_run, force_no_bold=True, force_bold=True)
        matched_any = True
        pos = match.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], source_run, force_no_bold=base_no_bold)
    if not matched_any and not text:
        _add_run(paragraph, "", source_run, force_no_bold=base_no_bold)


def _set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    source_run=None,
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> None:
    source_run = source_run if source_run is not None else _first_run_or_none(paragraph)
    _clear_paragraph_keep_ppr(paragraph)
    if keywords:
        _add_text_with_keyword_bold(paragraph, text, source_run, keywords, base_no_bold=True)
    else:
        _add_run(paragraph, _plain_resume_text(text), source_run, force_no_bold=force_no_bold)


def _insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    like_paragraph: Paragraph | None = None,
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> Paragraph:
    like_paragraph = like_paragraph or paragraph
    new_p = OxmlElement("w:p")
    if like_paragraph._p.pPr is not None:
        new_p.append(deepcopy(like_paragraph._p.pPr))
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    source_run = _first_meaningful_run_or_none(like_paragraph)
    if keywords:
        _add_text_with_keyword_bold(new_paragraph, text, source_run, keywords, base_no_bold=True)
    else:
        _add_run(new_paragraph, _plain_resume_text(text), source_run, force_no_bold=force_no_bold)
    return new_paragraph


def _delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def _is_decorative_or_blank_paragraph(paragraph: Paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return True
    return bool(re.fullmatch(r"[_\-—–=]{5,}", text))


def _replace_paragraph_with_lines(
    paragraph: Paragraph,
    lines: list[str],
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> None:
    lines = [str(line or "").strip() for line in lines if str(line or "").strip()]
    if not lines:
        _set_paragraph_text(paragraph, "", force_no_bold=force_no_bold)
        return
    source_run = _first_meaningful_run_or_none(paragraph)
    _set_paragraph_text(paragraph, lines[0], source_run, keywords=keywords, force_no_bold=force_no_bold)
    cursor = paragraph
    for line in lines[1:]:
        cursor = _insert_paragraph_after(cursor, line, like_paragraph=paragraph, keywords=keywords, force_no_bold=force_no_bold)


def _normalized_heading_text(paragraph: Paragraph) -> str:
    return _clean_text(paragraph.text).lower().strip(":")


def _is_section_heading(paragraph: Paragraph) -> bool:
    text = _normalized_heading_text(paragraph)
    if text in ALL_SECTION_TITLES:
        return True
    style_name = str(getattr(paragraph.style, "name", "") or "").lower()
    return bool(text) and "heading" in style_name and len(text.split()) <= 4


def _section_name(paragraph: Paragraph) -> str:
    text = _normalized_heading_text(paragraph)
    for key, aliases in SECTION_ALIASES.items():
        if text in aliases:
            return key
    return ""


def _find_section_range(paragraphs: list[Paragraph], section_key: str) -> tuple[int, int] | None:
    start = -1
    for idx, paragraph in enumerate(paragraphs):
        if _section_name(paragraph) == section_key:
            start = idx
            break
    if start < 0:
        return None
    end = len(paragraphs)
    for idx in range(start + 1, len(paragraphs)):
        if _is_section_heading(paragraphs[idx]):
            end = idx
            break
    return start, end


def _paragraphs_between_sections(doc: Document, section_key: str) -> list[Paragraph]:
    paragraphs = _all_body_paragraphs(doc)
    section_range = _find_section_range(paragraphs, section_key)
    if not section_range:
        return []
    start, end = section_range
    return paragraphs[start + 1:end]


def _replace_section_body(
    doc: Document,
    section_key: str,
    lines: list[str],
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> bool:
    paragraphs = _all_body_paragraphs(doc)
    section_range = _find_section_range(paragraphs, section_key)
    if not section_range:
        return False
    start, end = section_range
    body = paragraphs[start + 1:end]
    content_body = [p for p in body if _clean_text(p.text) and not _is_decorative_or_blank_paragraph(p)]
    if content_body:
        anchor = content_body[0]
        # Preserve blank/decorative paragraphs; remove only old content body.
        for p in content_body[1:]:
            _delete_paragraph(p)
        _replace_paragraph_with_lines(anchor, lines, keywords=keywords, force_no_bold=force_no_bold)
    else:
        anchor = paragraphs[start]
        cursor = anchor
        for line in [line for line in lines if str(line).strip()]:
            cursor = _insert_paragraph_after(cursor, line, like_paragraph=anchor, keywords=keywords, force_no_bold=force_no_bold)
    return True


def _replace_placeholders(
    doc: Document,
    placeholders: set[str],
    lines: list[str],
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> bool:
    changed = False
    lowered = {item.lower() for item in placeholders}
    for paragraph in _all_story_paragraphs(doc):
        text = _clean_text(paragraph.text).lower()
        if text in lowered:
            _replace_paragraph_with_lines(paragraph, lines, keywords=keywords, force_no_bold=force_no_bold)
            changed = True
    return changed


def _copy_paragraph_text_with_replacements(paragraph: Paragraph, replacements: dict[str, str]) -> bool:
    if not replacements:
        return False
    cleaned = {marker: _plain_resume_text(value) for marker, value in replacements.items()}
    original = paragraph.text or ""
    desired = original
    for marker, value in cleaned.items():
        desired = re.sub(re.escape(marker), value, desired, flags=re.IGNORECASE)
    if desired == original:
        return False

    changed_in_runs = False
    for run in paragraph.runs:
        run_text = run.text or ""
        new_text = run_text
        for marker, value in cleaned.items():
            new_text = re.sub(re.escape(marker), value, new_text, flags=re.IGNORECASE)
        if new_text != run_text:
            run.text = _preserve_line_text(new_text)
            changed_in_runs = True
    if changed_in_runs and (paragraph.text or "") == desired:
        return True

    _set_paragraph_text(paragraph, desired, _first_meaningful_run_or_none(paragraph))
    return True


def _replace_inline_placeholders(doc: Document, replacements: dict[str, str]) -> bool:
    changed = False
    for paragraph in _all_story_paragraphs(doc):
        if _copy_paragraph_text_with_replacements(paragraph, replacements):
            changed = True
    return changed


def _set_keep_lines(paragraph: Paragraph) -> None:
    try:
        paragraph.paragraph_format.keep_together = True
    except Exception:
        pass
    try:
        ppr = paragraph._p.get_or_add_pPr()
        if ppr.find(qn("w:keepLines")) is None:
            ppr.append(OxmlElement("w:keepLines"))
    except Exception:
        pass


def _remove_existing_tabs(ppr) -> None:
    try:
        tabs = ppr.find(qn("w:tabs"))
        if tabs is not None:
            ppr.remove(tabs)
    except Exception:
        pass


def _content_width_twips(paragraph: Paragraph) -> int:
    try:
        section = paragraph.part.document.sections[0]
        return int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)
    except Exception:
        return 9000


def _set_role_line_right_tab(paragraph: Paragraph) -> None:
    """Align right-side date/meta with a right tab instead of many spaces.

    Long generated roles cannot be kept stable with raw spaces in DOCX/PDF.
    A right tab preserves the line visually and prevents WPS/Word from expanding
    100+ spaces during PDF export.
    """
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass
    try:
        ppr = paragraph._p.get_or_add_pPr()
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            ppr.remove(jc)
        _remove_existing_tabs(ppr)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(_content_width_twips(paragraph)))
        tabs.append(tab)
        ppr.append(tabs)
    except Exception:
        pass


def _append_tab(paragraph: Paragraph, source_run=None) -> None:
    run = paragraph.add_run()
    _copy_run_format(source_run, run)
    try:
        run.add_tab()
    except Exception:
        run.text = "\t"


def _split_role_line(original: str):
    text = str(original or "").replace("\r", "").replace("\n", " ")
    match = _ROLE_MARKER_PATTERN.search(text)
    if not match:
        return None
    before = text[:match.start()]
    after_marker = text[match.end():]
    pipe_offset = after_marker.find("|")
    if pipe_offset < 0:
        return {
            "before": before,
            "role_marker": match.group(0),
            "separator": "",
            "right_text": after_marker,
            "after": "",
            "has_pipe": False,
        }
    bounded = after_marker[:pipe_offset]
    after = after_marker[pipe_offset + 1:]
    sep_match = re.match(r"([ \t\u00a0]*)(.*)", bounded, flags=re.DOTALL)
    separator = sep_match.group(1) if sep_match else ""
    right_text = (sep_match.group(2) if sep_match else bounded).strip()
    return {
        "before": before,
        "role_marker": match.group(0),
        "separator": separator,
        "right_text": right_text,
        "after": after,
        "has_pipe": True,
    }


def _replace_role_paragraph(paragraph: Paragraph, role_value: str) -> bool:
    original = paragraph.text or ""
    parts = _split_role_line(original)
    if not parts:
        return False
    role = _plain_resume_text(role_value)
    if not role:
        return False

    if not parts["has_pipe"]:
        replacements = {marker: role for marker in EXP_ROLE_PLACEHOLDERS}
        return _copy_paragraph_text_with_replacements(paragraph, replacements)

    before_run = _first_meaningful_run_or_none(paragraph)
    role_run = before_run
    right_run = before_run
    _clear_paragraph_keep_ppr(paragraph)
    _set_role_line_right_tab(paragraph)
    if parts["before"]:
        _add_run(paragraph, str(parts["before"]), before_run)
    _add_run(paragraph, role, role_run)
    if parts["right_text"]:
        _append_tab(paragraph, role_run)
        _add_run(paragraph, str(parts["right_text"]), right_run)
    if parts["after"]:
        _add_run(paragraph, str(parts["after"]).replace("|", ""), right_run)
    _set_keep_lines(paragraph)
    return True


def _replace_role_placeholders(doc: Document, resume: dict) -> bool:
    jobs = resume.get("work_history", []) or []
    role_values = [_plain_resume_text(job.get("role_title", "")) for job in jobs]
    role_values = [value for value in role_values if value]
    fallback_role = role_values[0] if role_values else _plain_resume_text(resume.get("headline", ""))
    if not fallback_role:
        return False
    role_index = 0
    changed = False
    for paragraph in _all_story_paragraphs(doc):
        if not _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
            continue
        role_value = role_values[min(role_index, len(role_values) - 1)] if role_values else fallback_role
        role_index += 1
        if _replace_role_paragraph(paragraph, role_value):
            changed = True
    return changed


def _parse_skill_groups(value) -> list[dict]:
    if isinstance(value, str):
        raw = value.strip()
        if raw:
            try:
                parsed = json.loads(raw)
                if isinstance(parsed, list):
                    value = parsed
            except Exception:
                value = []
    return value if isinstance(value, list) else []


def _skill_lines(resume: dict) -> list[str]:
    groups = _parse_skill_groups(resume.get("skill_groups") or [])
    lines: list[str] = []
    for group in groups:
        if not isinstance(group, dict):
            continue
        category = _plain_resume_text(group.get("category", ""))
        items = [_plain_resume_text(item) for item in group.get("items", []) if _plain_resume_text(item)]
        if category and items:
            lines.append(f"{category}: {', '.join(items)}")
    if lines:
        return lines
    skills = [_plain_resume_text(item) for item in resume.get("technical_skills", []) if _plain_resume_text(item)]
    return [", ".join(skills)] if skills else []


def _skill_groups_for_rendering(resume: dict) -> list[dict]:
    """Return parsed skill groups, or a single group from flat technical_skills."""
    groups = _parse_skill_groups(resume.get("skill_groups") or [])
    result = []
    for group in groups:
        if not isinstance(group, dict):
            continue
        category = _plain_resume_text(group.get("category", ""))
        items = [_plain_resume_text(item) for item in group.get("items", []) if _plain_resume_text(item)]
        if category and items:
            result.append({"category": category, "items": items})
    if result:
        return result
    skills = [_plain_resume_text(item) for item in resume.get("technical_skills", []) if _plain_resume_text(item)]
    if skills:
        return [{"category": "", "items": skills}]
    return []


def _write_skill_groups_bold_categories(
    anchor_paragraph: Paragraph,
    groups: list[dict],
    source_run=None,
) -> Paragraph:
    """Write skill groups so that category names are bold and skill items are normal weight.

    Returns the last paragraph written so callers can continue inserting after it.
    Each group gets its own paragraph: [BOLD category][normal ": item1, item2, ..."]
    """
    if not groups:
        return anchor_paragraph

    source_run = source_run or _first_meaningful_run_or_none(anchor_paragraph)

    def _write_group_to_paragraph(paragraph: Paragraph, group: dict) -> None:
        _clear_paragraph_keep_ppr(paragraph)
        category = group["category"]
        items_text = ", ".join(group["items"])
        if category:
            bold_run = paragraph.add_run(category)
            if source_run is not None:
                _copy_run_format(bold_run, source_run)
            _force_run_bold(bold_run)
            sep_run = paragraph.add_run(f"   {items_text}")
            if source_run is not None:
                _copy_run_format(sep_run, source_run)
            _force_run_not_bold(sep_run)
        else:
            plain_run = paragraph.add_run(items_text)
            if source_run is not None:
                _copy_run_format(plain_run, source_run)
            _force_run_not_bold(plain_run)

    _write_group_to_paragraph(anchor_paragraph, groups[0])
    cursor = anchor_paragraph
    for group in groups[1:]:
        new_p = OxmlElement("w:p")
        if anchor_paragraph._p.pPr is not None:
            new_p.append(deepcopy(anchor_paragraph._p.pPr))
        cursor._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, anchor_paragraph._parent)
        _write_group_to_paragraph(new_paragraph, group)
        cursor = new_paragraph
    return cursor


def _copy_run_format(target_run, source_run) -> None:
    """Copy font name and size from source_run to target_run without changing bold state."""
    try:
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
    except Exception:
        pass


def _insert_page_break_after(paragraph: Paragraph) -> Paragraph:
    """Insert a Word page break paragraph immediately after the given paragraph."""
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    new_br = OxmlElement("w:br")
    new_br.set(qn("w:type"), "page")
    new_r.append(new_br)
    new_p.append(new_r)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _paragraph_has_tab(paragraph: Paragraph) -> bool:
    return "\t" in (paragraph.text or "")


def _paragraph_has_numbering(paragraph: Paragraph) -> bool:
    try:
        ppr = paragraph._p.pPr
        return bool(ppr is not None and ppr.numPr is not None)
    except Exception:
        return False


def _looks_like_job_meta_line(paragraph: Paragraph) -> bool:
    """Protect company/location/duration rows from bullet replacement.

    Marvin-style resumes sometimes store company/location/duration as a
    numbered paragraph, so a pure numbering test misclassifies it as a bullet.
    These rows usually contain a company/location separator and a date range,
    and are often bold. They must stay unchanged unless explicit placeholders
    are present.
    """
    text = _clean_text(paragraph.text)
    if not text:
        return False
    if _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
        return False
    has_date = bool(_DATE_RE.search(text))
    has_company_separator = " | " in text or "|" in text
    has_many_spaces = bool(re.search(r"\S\s{8,}\S", paragraph.text or ""))
    if has_date and (has_company_separator or has_many_spaces):
        return True
    # Strong fallback for rows such as "Company Location  Jan 2020 – Present".
    if has_date and len(text.split()) <= 14 and not text.endswith("."):
        return True
    return False


def _looks_like_role_title_line(paragraph: Paragraph) -> bool:
    text = _clean_text(paragraph.text)
    if not text:
        return False
    if _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
        return False
    if _DATE_RE.search(text) or "|" in text:
        return False
    if text.endswith("."):
        return False
    words = text.split()
    if len(words) > 8:
        return False
    role_words = {"engineer", "developer", "lead", "manager", "architect", "analyst", "consultant", "specialist", "intern", "director"}
    return any(word.strip("(),/-").lower() in role_words for word in words)


def _paragraph_looks_like_bullet_content(paragraph: Paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return False
    if _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
        return False
    if _looks_like_job_meta_line(paragraph) or _looks_like_role_title_line(paragraph):
        return False
    if _paragraph_has_tab(paragraph):
        return False
    if text.endswith(":") and len(text.split()) <= 5:
        return False
    style_name = str(getattr(paragraph.style, "name", "") or "").lower()
    if "heading" in style_name:
        return False
    if text.startswith(("•", "-", "*", "‣", "◦")):
        return True
    if "bullet" in style_name:
        return True
    if "list" in style_name and len(text.split()) >= 6 and text.endswith((".", ";")):
        return True
    if _paragraph_has_numbering(paragraph) and len(text.split()) >= 6 and text.endswith((".", ";")):
        return True
    return False


def _replace_experience_bullets_and_titles(doc: Document, resume: dict) -> bool:
    body = _paragraphs_between_sections(doc, "experience")
    if not body:
        return False
    jobs = resume.get("work_history", []) or []
    if not jobs:
        return False
    changed = False
    keywords = _technical_skill_keywords(resume)

    bullet_groups: list[list[Paragraph]] = []
    current_group: list[Paragraph] = []
    for paragraph in body:
        if _paragraph_looks_like_bullet_content(paragraph):
            current_group.append(paragraph)
        else:
            if current_group:
                bullet_groups.append(current_group)
                current_group = []
    if current_group:
        bullet_groups.append(current_group)

    for idx, group in enumerate(bullet_groups):
        if idx >= len(jobs) or not group:
            continue
        bullets = [_plain_resume_text(bullet) for bullet in jobs[idx].get("bullets", []) if _plain_resume_text(bullet)]
        if not bullets:
            continue

        first = group[0]
        # Delete extra old bullets only inside this real bullet group; company,
        # role, duration, and location rows are not in this group.
        for old_paragraph in group[len(bullets):]:
            _delete_paragraph(old_paragraph)

        active_group = group[:len(bullets)]
        for paragraph, bullet in zip(active_group, bullets[:len(active_group)]):
            _set_paragraph_text(paragraph, bullet, keywords=keywords)
            changed = True

        cursor = active_group[-1] if active_group else first
        for bullet in bullets[len(active_group):]:
            cursor = _insert_paragraph_after(cursor, bullet, like_paragraph=first, keywords=keywords)
            changed = True

    return changed


def _replace_skills_with_bold_categories(doc: Document, resume: dict) -> bool:
    """Replace the skills section, rendering category names bold and skill items normal.

    Returns True if the section was found and updated.
    """
    skill_groups = _skill_groups_for_rendering(resume)
    if not skill_groups:
        return False

    paragraphs = _all_body_paragraphs(doc)
    # Try placeholder paragraphs first
    lowered_placeholders = {item.lower() for item in SKILL_PLACEHOLDERS}
    for paragraph in _all_story_paragraphs(doc):
        if _clean_text(paragraph.text).lower() in lowered_placeholders:
            _write_skill_groups_bold_categories(paragraph, skill_groups)
            return True

    # Fall back to locating the skills section
    section_range = _find_section_range(paragraphs, "skills")
    if not section_range:
        return False
    start, end = section_range
    body = paragraphs[start + 1:end]
    content_body = [p for p in body if _clean_text(p.text) and not _is_decorative_or_blank_paragraph(p)]
    if content_body:
        anchor = content_body[0]
        for p in content_body[1:]:
            _delete_paragraph(p)
        _write_skill_groups_bold_categories(anchor, skill_groups)
    else:
        anchor = paragraphs[start]
        source_run = _first_meaningful_run_or_none(anchor)
        cursor = anchor
        for group in skill_groups:
            new_p = OxmlElement("w:p")
            if anchor._p.pPr is not None:
                new_p.append(deepcopy(anchor._p.pPr))
            cursor._p.addnext(new_p)
            new_para = Paragraph(new_p, anchor._parent)
            _write_skill_groups_bold_categories(new_para, [group], source_run)
            cursor = new_para
    return True


def _trim_skills_to_fit(doc: Document, resume: dict, max_lines: int = 20) -> None:
    """Trim skill items per category until the skills section fits within max_lines.

    Each category row = 1 printed paragraph. We estimate its rendered line count
    at ~95 chars/line. If total skill lines exceed max_lines, we remove items from
    the end of the largest categories first until it fits.
    """
    _CHARS_PER_LINE = 95

    def _group_line_count(category: str, items: list[str]) -> int:
        line_text = f"{category}   {', '.join(items)}" if category else ', '.join(items)
        return max(1, -(-len(line_text) // _CHARS_PER_LINE))

    def _para_lines(p: Paragraph) -> int:
        text = (p.text or '').strip()
        if not text:
            return 1
        return max(1, -(-len(text) // _CHARS_PER_LINE))

    skill_groups = _skill_groups_for_rendering(resume)
    if not skill_groups:
        return

    total = sum(_group_line_count(g['category'], g['items']) for g in skill_groups)
    if total <= max_lines:
        return

    # Trim items from the end of the largest groups first, one item at a time
    groups = [{'category': g['category'], 'items': list(g['items'])} for g in skill_groups]
    while total > max_lines:
        # Find group with most items that still has >1 item
        target = max((g for g in groups if len(g['items']) > 1), key=lambda g: len(g['items']), default=None)
        if target is None:
            break
        old_lines = _group_line_count(target['category'], target['items'])
        target['items'].pop()
        new_lines = _group_line_count(target['category'], target['items'])
        total -= (old_lines - new_lines)

    # Write trimmed groups back into the doc
    paragraphs = _all_body_paragraphs(doc)
    section_range = _find_section_range(paragraphs, "skills")
    if not section_range:
        return
    start, end = section_range
    body = paragraphs[start + 1:end]
    content_body = [p for p in body if _clean_text(p.text) and not _is_decorative_or_blank_paragraph(p)]
    if not content_body:
        return
    anchor = content_body[0]
    for p in content_body[1:]:
        _delete_paragraph(p)
    _write_skill_groups_bold_categories(anchor, groups)


def _trim_skills_spear2(doc: Document, resume: dict) -> None:
    """spear-2 variant: measure everything else on page 1 (header, summary, education)
    and trim skills so that summary + skills + education all fit within one page.

    Page 1 layout (top to bottom): header block → SUMMARY heading → summary →
    TECHNICAL SKILLS heading → skills body → EDUCATION heading → education body.
    We count lines for everything except the skills body, subtract from the
    page budget, and that gives the skills budget.
    """
    _CHARS_PER_LINE = 95
    _LINES_PER_PAGE = 43
    # Each section heading on page 1 has a top border + w:spacing after=80 which
    # adds extra vertical space beyond what character counting captures.
    # 3 headings on page 1: SUMMARY, SKILLS, EDUCATION → 3 lines of overhead.
    _HEADING_OVERHEAD = 3

    def _para_lines(p: Paragraph) -> int:
        text = (p.text or '').strip()
        # Blank paragraphs still consume vertical space in the printed doc.
        if not text:
            return 1
        return max(1, -(-len(text) // _CHARS_PER_LINE))

    def _group_line_count(category: str, items: list[str]) -> int:
        line_text = f"{category}   {', '.join(items)}" if category else ', '.join(items)
        return max(1, -(-len(line_text) // _CHARS_PER_LINE))

    paragraphs = _all_body_paragraphs(doc)
    skills_range = _find_section_range(paragraphs, "skills")
    edu_range = _find_section_range(paragraphs, "education")
    exp_range = _find_section_range(paragraphs, "experience")

    if not skills_range:
        return

    skills_start, skills_end = skills_range

    # Lines consumed by everything before the skills body (header + summary section + SKILLS heading)
    pre_skills_lines = sum(_para_lines(p) for p in paragraphs[:skills_start + 1])

    # Lines consumed by education section (heading + body), stopping at experience
    post_skills_lines = 0
    if edu_range:
        edu_start, _ = edu_range
        edu_end_actual = exp_range[0] if exp_range else len(paragraphs)
        post_skills_lines = sum(_para_lines(p) for p in paragraphs[edu_start:edu_end_actual])

    fixed_lines = pre_skills_lines + post_skills_lines + _HEADING_OVERHEAD
    skills_budget = _LINES_PER_PAGE - fixed_lines

    # Minimum of 5 lines so we always show something
    skills_budget = max(5, skills_budget)

    skill_groups = _skill_groups_for_rendering(resume)
    if not skill_groups:
        return

    current_total = sum(_group_line_count(g['category'], g['items']) for g in skill_groups)
    if current_total <= skills_budget:
        return

    groups = [{'category': g['category'], 'items': list(g['items'])} for g in skill_groups]
    while current_total > skills_budget:
        target = max((g for g in groups if len(g['items']) > 1), key=lambda g: len(g['items']), default=None)
        if target is None:
            break
        old_lines = _group_line_count(target['category'], target['items'])
        target['items'].pop()
        new_lines = _group_line_count(target['category'], target['items'])
        current_total -= (old_lines - new_lines)

    # Write trimmed groups back into the doc
    # Re-query paragraphs since nothing changed structurally yet
    paragraphs = _all_body_paragraphs(doc)
    skills_range = _find_section_range(paragraphs, "skills")
    if not skills_range:
        return
    skills_start, skills_end = skills_range
    body = paragraphs[skills_start + 1:skills_end]
    content_body = [p for p in body if _clean_text(p.text) and not _is_decorative_or_blank_paragraph(p)]
    if not content_body:
        return
    anchor = content_body[0]
    for p in content_body[1:]:
        _delete_paragraph(p)
    _write_skill_groups_bold_categories(anchor, groups)


def _strip_paragraph_borders(doc: Document) -> None:
    """Remove any w:pBdr from every paragraph (template cleanup before content is written)."""
    for paragraph in doc.paragraphs:
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is not None:
            pPr.remove(pBdr)


def _is_section_heading_paragraph(p, resume_template: str) -> bool:
    """Return True if this paragraph element is a section heading (not the Full Name).

    spear-1: headings have runs with w:color=00B0F0.
    spear-2: headings are Heading 3 with w:color=000000 AND font size ≤ 28 half-pts
             (the Full Name is 72 half-pts and is excluded this way).
    """
    if resume_template == 'spear-2':
        # Must have a black-colored run
        has_black_run = False
        for rPr in p.iter(qn("w:rPr")):
            color_el = rPr.find(qn("w:color"))
            if color_el is not None:
                val = (color_el.get(qn("w:val")) or "").upper()
                if val == "000000":
                    has_black_run = True
                    break
        if not has_black_run:
            return False
        # Exclude the Full Name by font size (name = 72 half-pts, headings ≤ 28)
        for rPr in p.iter(qn("w:rPr")):
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                try:
                    if int(sz.get(qn("w:val"), "0")) > 28:
                        return False
                except ValueError:
                    pass
        return True
    else:
        # spear-1: blue run color marks section headings
        for rPr in p.iter(qn("w:rPr")):
            color_el = rPr.find(qn("w:color"))
            if color_el is not None:
                val = (color_el.get(qn("w:val")) or "").upper()
                if val == "00B0F0":
                    return True
        return False


def _inject_section_heading_borders(doc: Document, resume_template: str = 'spear-1') -> None:
    """Add a border to section headings only (called after all content is written).

    spear-1: blue (00B0F0) bottom border — line sits below the heading text.
    spear-2: black (000000) top border — line sits above the heading text.
    Full Name paragraph is always excluded. No page-break attributes added.
    """
    border_color = "000000" if resume_template == 'spear-2' else "00B0F0"
    border_side = "top" if resume_template == 'spear-2' else "bottom"

    for paragraph in doc.paragraphs:
        p = paragraph._p

        if not _is_section_heading_paragraph(p, resume_template):
            continue

        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)

        # remove any existing border first (safety)
        existing = pPr.find(qn("w:pBdr"))
        if existing is not None:
            pPr.remove(existing)

        new_pBdr = OxmlElement("w:pBdr")
        border_el = OxmlElement(f"w:{border_side}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "6")
        border_el.set(qn("w:space"), "1")
        border_el.set(qn("w:color"), border_color)
        new_pBdr.append(border_el)
        pPr.append(new_pBdr)

        # spear-2: add space after the heading so text below isn't flush against the border
        if resume_template == 'spear-2':
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.insert(0, spacing)
            spacing.set(qn("w:after"), "80")


def apply_resume_to_docx(docx_path: Path, resume: dict, resume_template: str = 'spear-1') -> None:
    doc = Document(str(docx_path))
    headline = _plain_resume_text(resume.get("headline", ""))
    summary = _plain_resume_text(resume.get("summary", ""))
    tech_keywords = _technical_skill_keywords(resume)
    _strip_paragraph_borders(doc)

    if headline:
        _replace_inline_placeholders(
            doc,
            {
                "___resume_title___": headline,
                "___headline___": headline,
                "__resume_title__": headline,
                "__headline__": headline,
            },
        )
    _replace_role_placeholders(doc, resume)

    if summary:
        if not _replace_placeholders(doc, SUMMARY_PLACEHOLDERS, [summary], keywords=tech_keywords):
            _replace_section_body(doc, "summary", [summary], keywords=tech_keywords)

    # Skills: category names bold, individual skills normal weight
    if not _replace_skills_with_bold_categories(doc, resume):
        skills = _skill_lines(resume)
        if skills:
            _replace_section_body(doc, "skills", skills, force_no_bold=True)

    # Trim skills to fit within their page budget (must run before page breaks are inserted)
    if resume_template == 'spear-2':
        _trim_skills_spear2(doc, resume)
    else:
        _trim_skills_to_fit(doc, resume, max_lines=20)

    experience_lines = []
    for job_index, job in enumerate(resume.get("work_history", []) or []):
        if job_index:
            experience_lines.append("")
        meta = " | ".join(
            item
            for item in [
                _plain_resume_text(job.get("company_name", "")),
                _plain_resume_text(job.get("duration", "")),
                _plain_resume_text(job.get("location", "")),
            ]
            if item
        )
        if meta:
            experience_lines.append(meta)
        role = _plain_resume_text(job.get("role_title", ""))
        if role:
            experience_lines.append(role)
        for bullet in job.get("bullets", []) or []:
            bullet_text = _plain_resume_text(bullet)
            if bullet_text:
                experience_lines.append(f"• {bullet_text}")

    if experience_lines:
        if not _replace_placeholders(doc, EXPERIENCE_PLACEHOLDERS, experience_lines, keywords=tech_keywords):
            _replace_experience_bullets_and_titles(doc, resume)

    # spear-2: forced page breaks — education end → new page, company 2 end → new page
    if resume_template == 'spear-2':
        _insert_spear2_page_breaks(doc, resume)

    # Section heading borders — style depends on template
    _inject_section_heading_borders(doc, resume_template=resume_template)

    doc.save(str(docx_path))


def _insert_experience_page_breaks(doc: Document, resume: dict) -> None:
    """Insert a page break after company 1's last bullet and after company 3's last bullet.

    This forces:  page 1 = header+summary+company1,  page 2 = company2+company3,  page 3 = company4+education+skills
    Only fires when there are at least 2 companies; only inserts break after company 3 when there are at least 4.
    """
    jobs = resume.get("work_history", []) or []
    if len(jobs) < 2:
        return

    paragraphs = _all_body_paragraphs(doc)

    # Find bullet groups inside the experience section
    exp_range = _find_section_range(paragraphs, "experience")
    if not exp_range:
        return
    exp_start, exp_end = exp_range
    exp_body = paragraphs[exp_start + 1:exp_end]

    bullet_group_last: list[Paragraph] = []
    current_group: list[Paragraph] = []
    for paragraph in exp_body:
        if _paragraph_looks_like_bullet_content(paragraph):
            current_group.append(paragraph)
        else:
            if current_group:
                bullet_group_last.append(current_group[-1])
                current_group = []
    if current_group:
        bullet_group_last.append(current_group[-1])

    # Insert page break after company 1 (index 0)
    if len(bullet_group_last) >= 1:
        _insert_page_break_after(bullet_group_last[0])

    # Insert page break after company 3 (index 2) — only when 4+ companies exist
    if len(jobs) >= 4 and len(bullet_group_last) >= 3:
        # Re-query paragraphs after the first insertion shifted offsets
        paragraphs2 = _all_body_paragraphs(doc)
        exp_range2 = _find_section_range(paragraphs2, "experience")
        if exp_range2:
            exp_start2, exp_end2 = exp_range2
            exp_body2 = paragraphs2[exp_start2 + 1:exp_end2]
            bullet_group_last2: list[Paragraph] = []
            current_group2: list[Paragraph] = []
            for paragraph in exp_body2:
                if _paragraph_looks_like_bullet_content(paragraph):
                    current_group2.append(paragraph)
                else:
                    if current_group2:
                        bullet_group_last2.append(current_group2[-1])
                        current_group2 = []
            if current_group2:
                bullet_group_last2.append(current_group2[-1])
            if len(bullet_group_last2) >= 3:
                _insert_page_break_after(bullet_group_last2[2])


def _insert_spear2_page_breaks(doc: Document, resume: dict) -> None:
    """spear-2 layout: page 1 = header+summary+skills+education, page 2 = companies 1+2, page 3 = companies 3+4.

    Inserts two page breaks:
      1. After the last paragraph of the Education section → Experience starts on page 2.
      2. After company 2's last bullet inside Experience → companies 3+4 start on page 3.
    """
    paragraphs = _all_body_paragraphs(doc)

    # --- Break 1: after Education section end ---
    edu_range = _find_section_range(paragraphs, "education")
    if edu_range:
        edu_start, edu_end = edu_range
        # Remove ALL empty paragraphs between the last education content and the
        # next section heading — they come from the template and would create a
        # blank gap at the top of page 2.
        empty_to_remove = []
        for idx in range(edu_end - 1, edu_start, -1):
            if not (paragraphs[idx].text or "").strip():
                empty_to_remove.append(paragraphs[idx])
            else:
                break
        for ep in empty_to_remove:
            ep._p.getparent().remove(ep._p)

        # Re-query after deletions
        paragraphs = _all_body_paragraphs(doc)
        edu_range = _find_section_range(paragraphs, "education")
        if edu_range:
            edu_start, edu_end = edu_range
            # Insert page break after the last non-empty education paragraph
            insert_after_idx = edu_end - 1
            while insert_after_idx > edu_start and not (paragraphs[insert_after_idx].text or "").strip():
                insert_after_idx -= 1
            _insert_page_break_after(paragraphs[insert_after_idx])

    # --- Break 2: after company 2's last bullet inside Experience ---
    # Re-query after first break shifted offsets
    paragraphs2 = _all_body_paragraphs(doc)
    exp_range2 = _find_section_range(paragraphs2, "experience")
    if not exp_range2:
        return
    exp_start2, exp_end2 = exp_range2
    exp_body2 = paragraphs2[exp_start2 + 1:exp_end2]

    # Collect bullet groups (one group = one company's bullets)
    bullet_groups: list[list[Paragraph]] = []
    current_group2: list[Paragraph] = []
    for paragraph in exp_body2:
        if _paragraph_looks_like_bullet_content(paragraph):
            current_group2.append(paragraph)
        else:
            if current_group2:
                bullet_groups.append(current_group2)
                current_group2 = []
    if current_group2:
        bullet_groups.append(current_group2)

    if len(bullet_groups) < 2:
        return

    # Trim company 2 bullets if they overflow page 2.
    # Page 2 holds companies 1+2. Estimate lines at ~95 chars/line, ~43 printable
    # lines per page. Each bullet = 2 lines (120-175 chars). Non-bullet rows = 1 line.
    _CHARS_PER_LINE = 95
    _LINES_PER_PAGE = 43

    def _estimated_lines(p: Paragraph) -> int:
        text = (p.text or '').strip()
        if not text:
            return 1
        return max(1, -(-len(text) // _CHARS_PER_LINE))  # ceiling division

    # Build id sets for fast identity checks
    co1_ids = {id(p) for p in bullet_groups[0]}
    co2_bullets = bullet_groups[1]
    co2_ids = {id(p) for p in co2_bullets}

    # Split exp_body2 into: company1 block (up to and including co1 last bullet),
    # company2 block (everything after co1 last bullet up to and including co2 last bullet)
    co1_last_id = id(bullet_groups[0][-1])
    co2_last_id = id(co2_bullets[-1]) if co2_bullets else None

    co1_block: list[Paragraph] = []
    co2_block: list[Paragraph] = []
    in_co2_block = False
    for p in exp_body2:
        if not in_co2_block:
            co1_block.append(p)
            if id(p) == co1_last_id:
                in_co2_block = True
        else:
            co2_block.append(p)
            if co2_last_id and id(p) == co2_last_id:
                break

    # page2 = exp heading + co1 block + co2 block
    page2_lines = 1  # PROFESSIONAL EXPERIENCE heading
    page2_lines += sum(_estimated_lines(p) for p in co1_block)
    page2_lines += sum(_estimated_lines(p) for p in co2_block)
    overflow_lines = page2_lines - _LINES_PER_PAGE

    # Drop bullets from the end of company 2 until it fits
    if overflow_lines > 0 and co2_bullets:
        bullets_to_drop: list[Paragraph] = []
        freed = 0
        for bp in reversed(co2_bullets):
            if freed >= overflow_lines:
                break
            freed += _estimated_lines(bp)
            bullets_to_drop.append(bp)
        for bp in bullets_to_drop:
            bp._p.getparent().remove(bp._p)

    # Re-query after potential removals
    paragraphs3 = _all_body_paragraphs(doc)
    exp_range3 = _find_section_range(paragraphs3, "experience")
    if not exp_range3:
        return
    exp_start3, exp_end3 = exp_range3
    exp_body3 = paragraphs3[exp_start3 + 1:exp_end3]

    bullet_group_last3: list[Paragraph] = []
    current_group3: list[Paragraph] = []
    for paragraph in exp_body3:
        if _paragraph_looks_like_bullet_content(paragraph):
            current_group3.append(paragraph)
        else:
            if current_group3:
                bullet_group_last3.append(current_group3[-1])
                current_group3 = []
    if current_group3:
        bullet_group_last3.append(current_group3[-1])

    if len(bullet_group_last3) >= 2:
        _insert_page_break_after(bullet_group_last3[1])


def find_soffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("soffice.exe"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/snap/bin/libreoffice",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def find_wps() -> str | None:
    candidates = [
        shutil.which("wps"),
        shutil.which("wps.exe"),
        r"C:\Program Files\WPS Office\office6\wps.exe",
        r"C:\Program Files (x86)\WPS Office\office6\wps.exe",
        r"C:\Program Files\Kingsoft\WPS Office\office6\wps.exe",
        r"C:\Program Files (x86)\Kingsoft\WPS Office\office6\wps.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def export_pdf_via_docx2pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    try:
        from docx2pdf import convert  # type: ignore
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "PDF created via docx2pdf"
        return False, "docx2pdf ran but no PDF file was created"
    except Exception as exc:
        return False, f"docx2pdf failed: {exc!r}"


def export_pdf_via_word(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    if sys.platform != "win32":
        return False, "Word COM export is only supported on Windows"
    word = None
    doc = None
    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore
        pythoncom.CoInitialize()
        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass
        doc = word.Documents.Open(str(docx_path), ReadOnly=True, AddToRecentFiles=False, ConfirmConversions=False)
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            CreateBookmarks=1,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "PDF created via Microsoft Word"
        return False, "Word export ran but no PDF file was created"
    except Exception as exc:
        return False, f"Word export failed: {exc!r}"
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            import pythoncom  # type: ignore
            pythoncom.CoUninitialize()
        except Exception:
            pass


def export_pdf_via_libreoffice(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    soffice = find_soffice()
    if not soffice:
        return False, "LibreOffice not found"
    temp_dir = Path(tempfile.mkdtemp(prefix="lo_pdf_"))
    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(temp_dir), str(docx_path)],
            capture_output=True,
            text=True,
            check=False,
            timeout=90,
        )
        generated = temp_dir / f"{docx_path.stem}.pdf"
        if generated.exists() and generated.stat().st_size > 0:
            pdf_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(generated, pdf_path)
            return True, "PDF created via LibreOffice"
        message = result.stderr.strip() or result.stdout.strip() or "LibreOffice conversion failed"
        return False, f"LibreOffice export failed: {message}"
    except Exception as exc:
        return False, f"LibreOffice export failed: {exc!r}"
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def export_pdf_via_wps_custom(docx_path: Path, pdf_path: Path, pdf_cfg: dict) -> tuple[bool, str]:
    command_template = str(pdf_cfg.get("wps_pdf_command", "") or "").strip()
    if not command_template:
        if find_wps():
            return False, "WPS found, but no WPS custom PDF command is configured"
        return False, "WPS not found and no WPS custom PDF command is configured"
    command = command_template.format(input=str(docx_path), output=str(pdf_path))
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True, check=False)
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "PDF created via WPS custom command"
        message = result.stderr.strip() or result.stdout.strip() or "WPS custom command failed"
        return False, f"WPS export failed: {message}"
    except Exception as exc:
        return False, f"WPS export failed: {exc!r}"


def export_pdf(docx_path: Path, pdf_path: Path, pdf_cfg: dict | None = None) -> tuple[bool, str]:
    pdf_cfg = pdf_cfg or {}
    order = pdf_cfg.get("backend_order")
    if isinstance(order, str):
        order = [item.strip() for item in order.split(",") if item.strip()]
    if not isinstance(order, list) or not order:
        order = ["docx2pdf", "word", "libreoffice", "wps_custom"]
    backend_map = {
        "docx2pdf": lambda: export_pdf_via_docx2pdf(docx_path, pdf_path),
        "word": lambda: export_pdf_via_word(docx_path, pdf_path),
        "libreoffice": lambda: export_pdf_via_libreoffice(docx_path, pdf_path),
        "wps_custom": lambda: export_pdf_via_wps_custom(docx_path, pdf_path, pdf_cfg),
    }
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    messages: list[str] = []
    for backend in order:
        key = str(backend or "").strip().lower()
        fn = backend_map.get(key)
        if not fn:
            messages.append(f"{backend}: unknown backend")
            continue
        ok, message = fn()
        messages.append(f"{backend}: {message}")
        if ok:
            return True, message
    return False, " | ".join(messages)


def pdf_backend_status(pdf_cfg: dict | None = None) -> list[str]:
    pdf_cfg = pdf_cfg or {}
    lines = []
    try:
        from docx2pdf import convert  # noqa: F401
        lines.append("docx2pdf: OK - package available")
    except Exception as exc:
        lines.append(f"docx2pdf: NO - {exc!r}")
    lines.append("word: OK - Windows only; checked during export" if sys.platform == "win32" else "word: NO - not Windows")
    soffice = find_soffice()
    lines.append(f"libreoffice: OK - {soffice}" if soffice else "libreoffice: NO - not found")
    wps = find_wps()
    if str(pdf_cfg.get("wps_pdf_command", "") or "").strip():
        lines.append("wps_custom: OK - command configured")
    elif wps:
        lines.append(f"wps_custom: NO - WPS found at {wps}, but command is not configured")
    else:
        lines.append("wps_custom: NO - WPS not found")
    return lines


def _uploaded_resume_path(profile: dict) -> Path | None:
    upload = profile.get("uploaded_resume") if isinstance(profile.get("uploaded_resume"), dict) else {}
    candidates = [
        str(upload.get("path", "") or "").strip(),
        str(upload.get("storage_path", "") or "").strip(),
        str(upload.get("relative_path", "") or "").strip(),
    ]
    for value in candidates:
        if not value:
            continue
        path = Path(value).expanduser()
        if path.exists():
            return path
    return None


def build_pdf_preview_html(pdf_bytes: bytes, message: str = "") -> str:
    if not pdf_bytes:
        return f"""
        <div style='font-family:Arial,sans-serif;padding:20px;border:1px solid #fca5a5;border-radius:12px;background:#fff1f2;color:#7f1d1d;'>
          <h3 style='margin-top:0;'>PDF preview is unavailable</h3>
          <p>{html.escape(message or 'The PDF exporter did not return a PDF file.')}</p>
        </div>
        """
    encoded = base64.b64encode(pdf_bytes).decode("ascii")
    return f"""
    <div style='font-family:Arial,sans-serif;'>
      <iframe title='Resume PDF preview' src='data:application/pdf;base64,{encoded}' style='width:100%;height:1120px;border:1px solid #e5e7eb;border-radius:12px;background:#fff;'></iframe>
      <p style='font-size:12px;color:#64748b;margin-top:8px;'>Read-only PDF preview generated from the uploaded DOCX style. {html.escape(message)}</p>
    </div>
    """


def build_docx_style_pdf_bundle(resume: dict, profile: dict, output_dir: Path | str, pdf_cfg: dict | None = None) -> dict[str, bytes | str]:
    source_docx = _uploaded_resume_path(profile)
    if not source_docx:
        raise FileNotFoundError("no resume so must upload resume")
    temp_dir = Path(tempfile.mkdtemp(prefix="tailorresume_docx_pdf_"))
    try:
        working_docx = temp_dir / "styled_resume.docx"
        pdf_path = temp_dir / "styled_resume.pdf"
        shutil.copy2(source_docx, working_docx)
        resume_template = str(profile.get('resume_template') or 'spear-1').strip() or 'spear-1'
        apply_resume_to_docx(working_docx, resume, resume_template=resume_template)
        ok, message = export_pdf(working_docx, pdf_path, pdf_cfg or {})
        pdf_bytes = pdf_path.read_bytes() if ok and pdf_path.exists() else b""
        docx_bytes = working_docx.read_bytes()
        return {
            "pdf": pdf_bytes,
            "html": build_pdf_preview_html(pdf_bytes, message),
            "markdown": "",
            "docx": docx_bytes,
            "pdf_message": message,
        }
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def build_docx_template_pdf_bundle(profile: dict, output_dir: Path | str, pdf_cfg: dict | None = None) -> dict[str, bytes | str]:
    source_docx = _uploaded_resume_path(profile)
    if not source_docx:
        raise FileNotFoundError("no resume so must upload resume")
    temp_dir = Path(tempfile.mkdtemp(prefix="tailorresume_docx_template_pdf_"))
    try:
        working_docx = temp_dir / "resume_template.docx"
        pdf_path = temp_dir / "resume_template.pdf"
        shutil.copy2(source_docx, working_docx)
        ok, message = export_pdf(working_docx, pdf_path, pdf_cfg or {})
        pdf_bytes = pdf_path.read_bytes() if ok and pdf_path.exists() else b""
        docx_bytes = working_docx.read_bytes()
        return {
            "pdf": pdf_bytes,
            "html": build_pdf_preview_html(pdf_bytes, message),
            "markdown": "",
            "docx": docx_bytes,
            "pdf_message": message,
        }
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
